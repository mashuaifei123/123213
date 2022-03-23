# coding: utf-8
import pandas as pd
import numpy as np
import cx_Oracle
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL  # 单元格对齐方式位置模块
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 段落,表格对齐方式位置模块
from docx.enum.text import WD_LINE_SPACING  # 设置段落的行间距
from docx.enum.section import WD_ORIENT  # 章节的方向 PORTRAIT 纵向　LANDSCAPE 横向
from docx.enum.section import WD_SECTION  # 增加新章节 NEW_PAGE 下一页
from datetime import datetime

out_path = r'K:\mashuaifei\watson\\'

watsonssqlserver = "10.10.92.50"
watsonsservice_name = "watpro"
watsonstns_ip = '10.10.92.50:1521/watpro'
watsonssqluser = "watson"
watsonssqlpw = "Abc123456"


def SampleNameSplits_def(values: str):
    '''判断SampleName长度并切割'''
    ii = values.split(' ')
    if len(ii) == 12:
        return ii[:5] + ii[7:]
    return ii


class Connect_WATPRO:
    '''连接到watpro准备查询数据 完成后记得调用close()关闭连接 '''

    def __init__(self):
        self.dbc = cx_Oracle.connect(watsonssqluser, watsonssqlpw, watsonstns_ip)

    def query(self, sql: str):
        query_df = pd.read_sql(sql, con=self.dbc)
        return query_df

    def close(self):
        return self.dbc.close()


def Connect_watpro(sql: str) -> pd.DataFrame:
    '''连接到watpro执行sql查询返回pd.DataFrame'''
    dbc = cx_Oracle.connect(watsonssqluser, watsonssqlpw, watsonstns_ip)
    query_df = pd.read_sql(sql, con=dbc)
    dbc.close()
    return query_df


class WatsonSQLstr:
    StudynameFindStudyid = """SELECT Study.StudyName, Study.StudyId FROM Study"""
    StudyidFindRunid = """SELECT AnaRunLink.RunId, Study.StudyId FROM AnaRunLink, Study WHERE AnaRunLink.StudyId = Study.StudyId  AND AnaRunLink.HostStudyId = 76  AND AnaRunLink.studyid = 76"""


def CRROLTF(valueslist):
    '''判断标准曲线的准确度满足条件
      （实测浓度与理论浓度的比值）应在 85.0 % ~ 115.0 %之间
      （定量下限的准确度应在 80.0 % ~ 120.0 %之间）
    '''
    LEVELNUMBER, RangeOfLine, ANALYTEFLAGPERCENT, C_Ratio_ROL = valueslist
    if LEVELNUMBER == 1:
        if 80 <= C_Ratio_ROL <= 120:
            return True
        else:
            return False
    else:
        if 85 <= C_Ratio_ROL <= 115:
            return True
        else:
            return False
        pass
    return '标准曲线不在范围内'  # 这是一个出错情况


def CRROLTF_QC(valueslist):
    '''判断标准曲线的准确度满足条件
      （实测浓度与理论浓度的比值）应在 85.0 % ~ 115.0 %之间
    '''
    LEVELNUMBER, RangeOfLine, ANALYTEFLAGPERCENT, C_Ratio_ROL = valueslist
    if 85 <= C_Ratio_ROL <= 115:
        return True
    else:
        return False


def CRROLTF_20_15(valueslist):
    '''判断标准曲线的准确度满足条件
    0- 一个数它就变成了负数
    '''
    LEVELNUMBER, RangeOfLine, ANALYTEFLAGPERCENT, C_Ratio_ROL = valueslist
    if 0 - float(ANALYTEFLAGPERCENT) <= C_Ratio_ROL <= float(ANALYTEFLAGPERCENT):
        return True
    else:
        return False


class WatsonLoadSQL:
    '''加载watson数据库中数据的类'''

    def __init__(self, StudyName):
        self.__db = cx_Oracle.connect(watsonssqluser, watsonssqlpw, watsonstns_ip)  # 建立连接
        self.__cr = self.__db.cursor()  # 创建游标
        self.StudyName = StudyName  # 专题编号
        self.StudyID = self._Query_Studys_ID().get(self.StudyName)  # 专题ID
        self.StudyID_Runs = self.__Query_Studys_RunID()  # 专题Runs  [(76, 1), (76, 2), (76, 3), (76, 4), (76, 5), (76, 6)]
        self.Query_AnalyteDescription = self._Query_AnalyteDescription()
        self.StudyID_Runs_R2 = self.__Query_RSQUARED_RunID()  # 专题所有R2结果  [(76, 5, 0.996642799040122, 'OK'), (76, 6, 0.994627870765527, 'OK')]
        self.StudyID_RawData_All = pd.merge(self.__Query_RawData_All(), self._Query_AnalyteDescription(),
                                            how='left')  # 获得该专题所有run结果 DataFrame
        self.StudyID_CONCENTRATION = self.__Query_AssayID_CONCENTRATION()

    def _Query_Studys_ID(self):
        '''返回StudysID 'A2020010-K10-01', 76'''
        StudyName_IDSql = r"""SELECT
                            StudyName,
                            StudyId
                            FROM
                            Study
                            WHERE
                            STUDYNAME = '{}'
                            """.format(self.StudyName)
        self.__cr.execute(StudyName_IDSql)  # 执行语句
        rs = self.__cr.fetchall()  # 获取数据全部
        return {k: v for k, v in rs}

    def __Query_Studys_RunID(self):
        '''返回专题Runs  [(76, 1), (76, 2), (76, 3), (76, 4), (76, 5), (76, 6)]'''
        StudyID_RunIDSql = r"""SELECT
                                Study.StudyId,
                                AnaRunLink.RunId
                                FROM
                                AnaRunLink,
                                Study
                                WHERE
                                AnaRunLink.StudyId = Study.StudyId 
                                AND AnaRunLink.HostStudyId = {} 
                                AND AnaRunLink.studyid = {}
                                """.format(self.StudyID, self.StudyID)
        self.__cr.execute(StudyID_RunIDSql)  # 执行语句
        rs = self.__cr.fetchall()  # 获取数据全部
        return rs

    def __Query_RSQUARED_RunID(self):
        '''返回专题Runs R2 [(76, 5, 0.996642799040122, 'OK')'''
        StudyID_R2sql = r"""SELECT
                            STUDYID,
                            RUNID,
                            ANALYTEINDEX,
                            RSQUARED,
                            ACCEPTREJECTREASON
                            FROM
                            ANALYTICALRUNANALYTES
                            WHERE
                            StudyId = {}
                            """.format(self.StudyID)
        self.__cr.execute(StudyID_R2sql)  # 执行语句
        rs = self.__cr.fetchall()  # 获取数据全部
        return rs

    def _Query_AnalyteDescription(self):
        '''返回 AnalyteDescription'''
        sql = r"""SELECT DISTINCT
                    AssayAnalytes.AnalyteIndex,
                    GlobalAnalytes.AnalyteDescription,
                    ConcentrationUnits.ConcentrationUnits,
                    AssayAnalytes.InternalStandard
                    FROM
                    ConfigRegressionTypes,
                    Assay,
                    AssayAnalytes,
                    GlobalAnalytes,
                    ConcentrationUnits 
                    WHERE
                    AssayAnalytes.AnalyteID = GlobalAnalytes.GlobalAnalyteId 
                    AND AssayAnalytes.Studyid = {}
                    AND Assay.assayId = AssayAnalytes.AssayID 
                    AND AssayAnalytes.ConcUnitsId = ConcentrationUnits.ConcUnitsId ( + ) 
                    AND ConfigRegressionTypes.regressionId ( + ) = AssayAnalytes.regressionIdentifier
                """.format(self.StudyID)
        self.__cr.execute(sql)  # 执行语句
        rs = self.__cr.fetchall()  # 获取数据全部
        return pd.DataFrame(rs,
                            columns=['AnalyteIndex', 'AnalyteDescription', 'ConcentrationUnits', 'InternalStandard'])

    def __Query_AssayID_CONCENTRATION(self):
        '''返回标准曲线'''
        sql = r"""SELECT DISTINCT
                    Assay.assayDescription,
                    Assay.StudyId,
                    Assay.RunId,
                    AssayAnalyteKnown.ASSAYID,
                    AssayAnalyteKnown.KNOWNTYPE,
                    AssayAnalyteKnown.ANALYTEINDEX,
                    AssayAnalyteKnown.LEVELNUMBER,
                    AssayAnalyteKnown.CONCENTRATION,
                    AssayAnalyteKnown.ANALYTEFLAGPERCENT,
                    GlobalAnalytes.AnalyteDescription
                    FROM
                    AssayAnalyteKnown,
                    AssayAnalytes,
                    Assay,
                    GlobalAnalytes 
                    WHERE
                    AssayAnalyteKnown.AnalyteIndex = AssayAnalytes.AnalyteIndex 
                    AND Assay.assayId = AssayAnalyteKnown.AssayId
                    AND AssayAnalyteKnown.AssayId = AssayAnalytes.AssayID 
                    AND AssayAnalytes.AnalyteID = GlobalAnalytes.GlobalAnalyteId 
                    AND AssayAnalyteKnown.STUDYID = {} 
                    ORDER BY
                    AssayAnalyteKnown.ASSAYID,
                    AssayAnalyteKnown.KnownType,
                    AssayAnalyteKnown.AnalyteIndex,
                    AssayAnalyteKnown.LevelNumber""".format(self.StudyID)
        self.__cr.execute(sql)  # 执行语句
        rs = self.__cr.fetchall()  # 获取数据全部
        return pd.DataFrame(rs, columns=['assayDescription', 'StudyId', 'runId', 'ASSAYID', 'KNOWNTYPE', 'ANALYTEINDEX',
                                         'LEVELNUMBER', 'CONCENTRATION', 'ANALYTEFLAGPERCENT', 'AnalyteDescription'])

    def __Query_RawData_All(self):
        '''返回所有专题数据DataFrame'''
        ARARawDatasql = r"""SELECT
                            AnalyticalRunSample.runId,
                            AnaRunRawAnalytePeak.SampleName,
                            AnaRunRawAnalytePeak.AnalyteArea,
                            AnaRunRawAnalytePeak.AnalytePeakRetentionTime,
                            AnaRunRawAnalytePeak.InternalStandardArea,
                            AnaRunRawAnalytePeak.InternalStandardRetentionTime,
                            AnaRunRawAnalytePeak.Ratio,
                            AnalyticalRunSample.runSampleKind,
                            AnaRunAnalyteResults.Concentration,
                            AnaRunAnalyteResults.eliminatedFlag,
                            ConfigAnalyteRunStatus.AnaRegStatusDesc,
                            AnaRunRawAnalytePeak.AnalyteIndex 
                            FROM
                            AnalyticalRunAnalytes,
                            AnaRunAnalyteResults,
                            AnaRunRawAnalytePeak,
                            ConfigAnalyteRunStatus,
                            AnalyticalRunSample 
                            WHERE
                            AnaRunRawAnalytePeak.RunSampleSequenceNumber = AnalyticalRunSample.RunSampleSequenceNumber 
                            AND AnalyticalRunAnalytes.RunAnalyteRegressionStatus = ConfigAnalyteRunStatus.RunAnalyteRegressionStatus ( + ) 
                            AND AnalyticalRunSample.studyId = {}
                            AND AnaRunRawAnalytePeak.studyId = {}
                            AND AnaRunRawAnalytePeak.RunID = AnalyticalRunSample.RunID 
                            AND AnaRunRawAnalytePeak.StudyId = AnalyticalRunSample.StudyId 
                            AND AnaRunAnalyteResults.AnalyteIndex = AnaRunRawAnalytePeak.AnalyteIndex 
                            AND AnaRunAnalyteResults.RunSampleSequenceNumber = AnaRunRawAnalytePeak.RunSampleSequenceNumber 
                            AND AnaRunAnalyteResults.RunID = AnaRunRawAnalytePeak.RunID 
                            AND AnaRunAnalyteResults.StudyId = AnaRunRawAnalytePeak.StudyId 
                            AND AnalyticalRunAnalytes.AnalyteIndex = AnaRunAnalyteResults.AnalyteIndex 
                            AND AnalyticalRunAnalytes.RunID = AnaRunAnalyteResults.RunID 
                            AND AnalyticalRunAnalytes.StudyId = AnaRunAnalyteResults.StudyId 
                            ORDER BY
                            AnalyticalRunSample.runId,
                            AnalyticalRunSample.RunSampleOrderNumber,
                            AnaRunRawAnalytePeak.AnalyteIndex
                            """.format(self.StudyID, self.StudyID)
        self.__cr.execute(ARARawDatasql)  # 执行语句
        rs = self.__cr.fetchall()  # 获取数据全部
        ARARawDataDF = pd.DataFrame(rs,
                                    columns=['runId', 'SampleName', 'Area', 'RetTime', 'ISArea', 'ISRetTime', 'Ratio',
                                             'SampleKind', 'Concentration', 'Status', 'AnaRegStatusDesc',
                                             'AnalyteIndex'])
        return ARARawDataDF

    def __del__(self):
        self.__cr.close()  # 关闭游标
        self.__db.close()  # 断开连接


class SystemAdaptability():  # 系统适应性
    '''适应性'''

    def __init__(self, StudyName, StudyID_Runs_R2, StudyID_RawData_All, StudyID_CONCENTRATION, Query_AnalyteDescription,
                 AnalyteDescription, SDT_Range, QC_Range):
        self.AnalyteDescription = AnalyteDescription
        self.StudyName = StudyName
        self.SDT_Range = SDT_Range
        self.QC_Range = QC_Range
        self.Query_AnalyteDescription = Query_AnalyteDescription
        self.StudyID_Runs_R2 = StudyID_Runs_R2
        self.StudyID_RawData_All = StudyID_RawData_All
        self.StudyID_CONCENTRATION = StudyID_CONCENTRATION
        self.rawdata = self.StudyID_RawData_All[
            self.StudyID_RawData_All.loc[:, 'AnalyteDescription'].isin([self.AnalyteDescription])].copy()
        self.ARARawDataDF_c = self._get_ARARawDataDF().copy()
        # self.QC_max = self._get_QC()[2]

    def _get_ARARawDataDF(self):
        df = self._get_LOQ()[1].copy()
        df = df[df.loc[:, 'Status'].isin(['N'])]  # 筛选 被拒绝的.
        df = df[df.loc[:, 'AnaRegStatusDesc'].isin(['Accepted'])]  # 晒选 Accepted
        df = df[df.loc[:, 'Carryover_Area_TF'] == True]
        df = df[df.loc[:, 'Carryover_ISArea_TF'] == True]
        return df

    def _get_RunR2(self):
        '''返回run r2'''
        r2df = pd.DataFrame(self.StudyID_Runs_R2, columns=['StudyID', 'RunID', 'AnalyteIndex', 'R2', 'Astatus'])
        r2df = pd.merge(r2df, self.Query_AnalyteDescription)
        # r2df['R2TF'] = r2df['R2'].apply(lambda x: True if float(x) > 0.98 else False) # 求取真假 大于0.98
        r2df['R2TF'] = [True if float(x) > 0.98 and y == 'OK' else False for x, y in
                        r2df[['R2', 'Astatus']].values.tolist()]
        r2df = r2df[r2df.loc[:, 'AnalyteDescription'].isin([self.AnalyteDescription])]  # 筛选 self.AnalyteDescription
        return r2df

    def _get_SDT(self):
        '''返回标准曲线'''
        ARARawDataDF_SDT = self.ARARawDataDF_c[self.ARARawDataDF_c.loc[:, 'SampleKind'].isin(['STANDARD'])][
            ['runId', 'SampleName', 'Concentration', 'ISArea']].copy()  # 筛选 STD并选择需要的列
        ARARawDataDF_SDT['LEVELNUMBER'] = [int(x.split(' ')[3][-1]) for x in ARARawDataDF_SDT['SampleName']]
        CONCENTRATION = self.StudyID_CONCENTRATION[
            (self.StudyID_CONCENTRATION.loc[:, 'KNOWNTYPE'].isin(['STANDARD'])) & (
                self.StudyID_CONCENTRATION.loc[:, 'AnalyteDescription'].isin([self.AnalyteDescription]))]
        # 筛选下STANDARD 和 AnalyteDescrition
        ARARawDataDF_SDT = pd.merge(ARARawDataDF_SDT, CONCENTRATION, on=['runId', 'LEVELNUMBER'])
        runid_name_list = list(filter(None,
                                      [runidsname if set(daf['CONCENTRATION']) == set(self.SDT_Range) else None for
                                       runidsname, daf in ARARawDataDF_SDT.groupby('runId')]))
        # 筛许满足SDT标准取值的数据RUNID
        ARARawDataDF_SDT = ARARawDataDF_SDT[ARARawDataDF_SDT.loc[:, 'runId'].isin(runid_name_list)]
        ARARawDataDF_SDT['RangeOfLine'] = ARARawDataDF_SDT['CONCENTRATION']  # 新增加列从SpampleName 获取线性范围
        ARARawDataDF_SDT['C_Ratio_ROL'] = ARARawDataDF_SDT['Concentration'] / ARARawDataDF_SDT[
            'RangeOfLine'] * 100  # 新增加列C_Ratio_ROL 求比值
        ARARawDataDF_SDT['CRROL_T_F'] = list(map(CRROLTF, ARARawDataDF_SDT[
            ['LEVELNUMBER', 'RangeOfLine', 'ANALYTEFLAGPERCENT', 'C_Ratio_ROL']].values.tolist()))
        ARARawDataDF_SDT['C_subtraction_ROL_Ratio_ROL'] = (ARARawDataDF_SDT['Concentration'] - ARARawDataDF_SDT[
            'RangeOfLine']) / ARARawDataDF_SDT['RangeOfLine']
        ARARawDataDF_SDT['CSRRR_T_F'] = list(map(CRROLTF_20_15, ARARawDataDF_SDT[
            ['LEVELNUMBER', 'RangeOfLine', 'ANALYTEFLAGPERCENT', 'C_subtraction_ROL_Ratio_ROL']].values.tolist()))
        # 其他的分开标记最大最小 共计4 个数值 比值需要不包含 False
        ARARawDataDF_SDT_T = ARARawDataDF_SDT[ARARawDataDF_SDT['CRROL_T_F'] == True]
        ARARawDataDF_SDT_T = ARARawDataDF_SDT[ARARawDataDF_SDT['CSRRR_T_F'] == True]
        LLOQ1 = 1
        A10min = ARARawDataDF_SDT_T[ARARawDataDF_SDT_T['LEVELNUMBER'] == LLOQ1]['C_Ratio_ROL'].min()
        A10max = ARARawDataDF_SDT_T[ARARawDataDF_SDT_T['LEVELNUMBER'] == LLOQ1]['C_Ratio_ROL'].max()
        A0min = ARARawDataDF_SDT_T[ARARawDataDF_SDT_T['LEVELNUMBER'] != LLOQ1]['C_Ratio_ROL'].min()
        A0max = ARARawDataDF_SDT_T[ARARawDataDF_SDT_T['LEVELNUMBER'] != LLOQ1]['C_Ratio_ROL'].max()
        Aalldf = pd.DataFrame([[A10min, A10max], [A0min, A0max]], columns=['MIN %', 'MAX %'], index=['LLOQ', 'OTHERS'])
        ARARawDataDF_SDT = ARARawDataDF_SDT[
            ['runId', 'SampleName', 'Concentration', 'ISArea', 'LEVELNUMBER', 'assayDescription', 'KNOWNTYPE',
             'CONCENTRATION', 'ANALYTEFLAGPERCENT', 'AnalyteDescription', 'RangeOfLine', 'C_Ratio_ROL',
             'CRROL_T_F', 'C_subtraction_ROL_Ratio_ROL', 'CSRRR_T_F']]
        return ARARawDataDF_SDT, Aalldf

    def _get_QC(self):
        '''返回QC'''
        ARARawDataDF_QC = self.ARARawDataDF_c[self.ARARawDataDF_c.loc[:, 'SampleKind'].isin(['QC'])][
            ['runId', 'SampleName', 'Concentration', 'ISArea']].copy()  # 筛选 QC并选择需要的列
        if len(self.QC_Range) == 3:  # 传入的数据长度.
            QC_Range_D = {'LQC': 1, 'MQC': 2, 'HQC': 3}
        else:
            QC_Range_D = {'LLOQ': 1, 'LQC': 2, 'MQC': 3, 'HQC': 4}
        CONCENTRATION = self.StudyID_CONCENTRATION[(self.StudyID_CONCENTRATION.loc[:, 'KNOWNTYPE'].isin(['QC'])) & (
            self.StudyID_CONCENTRATION.loc[:, 'AnalyteDescription'].isin([self.AnalyteDescription]))]
        # CONCENTRATION.to_excel(r'\\fileserver\部门共享\ANA\WatsonOUT\CONCEN0.xlsx')
        ARARawDataDF_QC['LEVELNUMBER'] = [QC_Range_D.get(x.split(' ')[3]) for x in
                                          ARARawDataDF_QC['SampleName']]  # 新增加列从SpampleName 获取线性范围
        # ARARawDataDF_QC.to_excel(r'\\fileserver\部门共享\ANA\WatsonOUT\QC1.xlsx')
        ARARawDataDF_QC = pd.merge(ARARawDataDF_QC, CONCENTRATION, on=['runId', 'LEVELNUMBER'])
        runid_name_list = list(filter(None, [runidsname if set(daf['CONCENTRATION']) == set(self.QC_Range) else None for
                                             runidsname, daf in ARARawDataDF_QC.groupby('runId')]))
        # 筛许满足QC标准取值的数据RUNID
        ARARawDataDF_QC = ARARawDataDF_QC[ARARawDataDF_QC.loc[:, 'runId'].isin(runid_name_list)]
        ARARawDataDF_QC['RangeOfLine'] = ARARawDataDF_QC['CONCENTRATION']  # 新增加列从SpampleName 获取线性范围
        ARARawDataDF_QC['C_Ratio_ROL'] = ARARawDataDF_QC['Concentration'] / ARARawDataDF_QC[
            'RangeOfLine'] * 100  # 新增加列C_Ratio_ROL 求比值
        ARARawDataDF_QC['CRROL_T_F'] = list(map(CRROLTF_QC, ARARawDataDF_QC[
            ['LEVELNUMBER', 'RangeOfLine', 'ANALYTEFLAGPERCENT', 'C_Ratio_ROL']].values.tolist()))
        ARARawDataDF_QC['C_subtraction_ROL_Ratio_ROL'] = (ARARawDataDF_QC['Concentration'] - ARARawDataDF_QC[
            'RangeOfLine']) / ARARawDataDF_QC['RangeOfLine']  # 新增加列C_Ratio_ROL 求比值
        ARARawDataDF_QC['CSRRR_T_F'] = list(map(CRROLTF_20_15, ARARawDataDF_QC[
            ['LEVELNUMBER', 'RangeOfLine', 'ANALYTEFLAGPERCENT', 'C_subtraction_ROL_Ratio_ROL']].values.tolist()))
        # 超出范围的增加星号标识 原始数据和计算值都要 转化为保留两位小数的字符串
        QCalldf = pd.DataFrame([['{:.2f}'.format(
            ARARawDataDF_QC[ARARawDataDF_QC.loc[:, 'CRROL_T_F'] == True]['C_Ratio_ROL'].min()), '{:.2f}'.format(
            ARARawDataDF_QC[ARARawDataDF_QC.loc[:, 'CRROL_T_F'] == True]['C_Ratio_ROL'].max())]],
            columns=['MIN %', 'MAX %'], index=['QC'])
        ARARawDataDF_QC['Concentration'] = [f'{b:.2f}' if a else f'{b:.2f}*' for a, b in
                                            zip(ARARawDataDF_QC['CRROL_T_F'], ARARawDataDF_QC['Concentration'])]
        ARARawDataDF_QC['C_Ratio_ROL'] = [f'{b:.2f}' if a else f'{b:.2f}*' for a, b in
                                          zip(ARARawDataDF_QC['CRROL_T_F'], ARARawDataDF_QC['C_Ratio_ROL'])]
        ARARawDataDF_QC = ARARawDataDF_QC[
            ['runId', 'SampleName', 'Concentration', 'ISArea', 'LEVELNUMBER', 'assayDescription', 'KNOWNTYPE',
             'CONCENTRATION', 'ANALYTEFLAGPERCENT', 'AnalyteDescription', 'RangeOfLine', 'C_Ratio_ROL',
             'CRROL_T_F', 'C_subtraction_ROL_Ratio_ROL', 'CSRRR_T_F']]
        return ARARawDataDF_QC, QCalldf

    def _get_LOQ(self):
        '''LOQ残留'''
        # 先对初始的数据进行筛选
        ARARawDataDF_LOQ = self.rawdata[self.rawdata.loc[:, 'SampleKind'].isin(['Carryover', 'STANDARD'])][
            ['runId', 'SampleName', 'Area', 'RetTime', 'ISArea', 'ISRetTime']].copy()  # 筛选 LOQ并选择需要的列
        alldata = self.rawdata
        alldata = alldata.reset_index(drop=False)
        # 如果有数字就把数字比上STD 11 和 STD 21 的均值 Area 和 ISARea分开  - 中横线替代nan
        ARALOQ_list = []
        for name, dfs in ARARawDataDF_LOQ.groupby('runId'):
            ARALOQ_list.append([name, dfs[dfs['SampleName'].str.contains('STD1')]['Area'].mean(),
                                dfs[dfs['SampleName'].str.contains('STD1')]['ISArea'].mean()])
        ARALOQDFMEAN = pd.DataFrame(ARALOQ_list, columns=['runId', 'Area_mean', 'ISArea_mean', ])
        ARARawDataDF_Carry = self.rawdata[self.rawdata.loc[:, 'SampleKind'].isin(['Carryover'])][
            ['runId', 'SampleName', 'Area', 'ISArea', ]].copy()  # 筛选 LOQ并选择需要的列
        ARARawDataDF_Carry['index_o'] = ARARawDataDF_Carry.index
        Carry_mean_df = pd.merge(ARARawDataDF_Carry, ARALOQDFMEAN, how='left')
        Carry_mean_df = Carry_mean_df.set_index(['index_o'])
        Carry_mean_df['Area_Ratio'] = ['-' if pd.isna(x) else float(x) / y for x, y in
                                       zip(Carry_mean_df['Area'], Carry_mean_df['Area_mean'])]
        Carry_mean_df['ISArea_Ratio'] = ['-' if pd.isna(x) else float(x) / y for x, y in
                                         zip(Carry_mean_df['ISArea'], Carry_mean_df['ISArea_mean'])]
        Carry_mean_df['Area_Ratio_TF'] = [True if x == '-' else (False if float(x) > 0.2 else True) for x in
                                          Carry_mean_df['Area_Ratio']]
        Carry_mean_df['ISArea_Ratio_TF'] = [True if x == '-' else (False if float(x) > 0.05 else True) for x in
                                            Carry_mean_df['ISArea_Ratio']]
        Area_d_20 = set(Carry_mean_df[Carry_mean_df.loc[:, 'Area_Ratio_TF'] == False][
                            'runId'])  # Area_Ratio 分析物峰面积大于20%  筛选获得runid
        ISArea_d_5 = set(Carry_mean_df[Carry_mean_df.loc[:, 'ISArea_Ratio_TF'] == False][
                             'runId'])  # ISArea_Ratio 分析物峰面积大于5%  筛选获得runid
        Area_df_list = [pd.DataFrame(columns=['index', 'runId', 'SampleName', 'Area_Ratio_values'])]
        ISArea_df_list = [pd.DataFrame(columns=['index', 'runId', 'SampleName', 'ISArea_Ratio_values'])]
        # 获取当大于Area_d_20情况下的runid进行筛选计算
        if Area_d_20:
            for aread20 in Area_d_20:  # 这里是有残留的runid
                aread20_df = alldata[alldata.loc[:, 'runId'] == aread20].copy()
                areaudlist = []
                for index_upv in aread20_df[
                    aread20_df['SampleName'].str.contains('Carryover')].index:  # 找到carryover的index
                    areaudlist.append(aread20_df.loc[[index_upv], ['Area']].values[0] /
                                      aread20_df.loc[[index_upv - 1], ['Area']].values[0])  # 上一列
                areamax = max(areaudlist)[0]  # 求出最大值
                QCUNdf = aread20_df[aread20_df.loc[:, 'SampleKind'].isin(['QC', 'UNKNOWN'])].copy()
                areazlist = QCUNdf['Area'].copy().tolist()
                QCUNdf['Area_Ratio_values'] = [x * areamax / y if (x and y) else None for x, y in
                                               zip([0] + areazlist, areazlist + [0])][:-1]  # 用错位的方式计算
                # 计算比值用于判断是否符合
                Area_df_list.append(QCUNdf[QCUNdf.loc[:, 'SampleKind'].isin(['UNKNOWN'])][
                                        ['index', 'runId', 'SampleName', 'Area_Ratio_values']])
        # 获取当大于ISArea大于5%情况下的runid进行计算
        if ISArea_d_5:
            for isarea5 in ISArea_d_5:  # 这里是有残留的runid
                isarea5_df = alldata[alldata.loc[:, 'runId'] == isarea5].copy()
                isareaudlist = []
                for index_upv in isarea5_df[
                    isarea5_df['SampleName'].str.contains('Carryover')].index:  # 找到carryover的index
                    isareaudlist.append(isarea5_df.loc[[index_upv], ['ISArea']].values[0] /
                                        aread20_df.loc[[index_upv - 1], ['ISArea']].values[0])  # 上一列
                isareamax = max(isareaudlist)[0]  # 求出最大值
                QCUNdf = isarea5_df[isarea5_df.loc[:, 'SampleKind'].isin(['QC', 'UNKNOWN'])].copy()
                isareazlist = QCUNdf['ISArea'].copy().tolist()
                QCUNdf['ISArea_Ratio_values'] = [x * isareamax / y if (x and y) else None for x, y in
                                                 zip([0] + isareazlist, isareazlist + [0])][:-1]  # 用错位的方式计算
                # 计算比值用于判断是否符合
                ISArea_df_list.append(QCUNdf[QCUNdf.loc[:, 'SampleKind'].isin(['UNKNOWN'])][
                                          ['index', 'runId', 'SampleName', 'ISArea_Ratio_values']])
        Area_df_d_20all = pd.concat(Area_df_list)
        Area_df_d_20all = Area_df_d_20all.astype({'index': 'int64', 'runId': 'int64', 'SampleName': 'object'})
        ISArea_df_d_5all = pd.concat(ISArea_df_list)
        ISArea_df_d_5all = ISArea_df_d_5all.astype({'index': 'int64', 'runId': 'int64', 'SampleName': 'object'})
        Carry_result_df = pd.merge(alldata, Area_df_d_20all, how='left')
        Carry_result_df = pd.merge(Carry_result_df, ISArea_df_d_5all, how='left')
        Carry_result_df['Carryover_Area_TF'] = [True if pd.isna(x) else (False if x > 0.05 else True) for x in
                                                Carry_result_df['Area_Ratio_values']]
        Carry_result_df['Carryover_ISArea_TF'] = [True if pd.isna(x) else (False if x > 0.05 else True) for x in
                                                  Carry_result_df['ISArea_Ratio_values']]
        # Carry_result_df.to_excel(r"C:\Users\rock\Desktop\watson\2000.xlsx")
        return Carry_mean_df, Carry_result_df

    def _get_Diff(self):
        '''返回内标'''
        AnmailD_DF = self.ARARawDataDF_c[(self.ARARawDataDF_c.loc[:, 'SampleKind'] == 'UNKNOWN') & (
            self.ARARawDataDF_c.loc[:, 'AnalyteDescription'].isin(
                [self.AnalyteDescription]))].copy()  # 筛许SampleKind名称为UNKNOWN 确认动物编号.
        SampleNameSplits = list(map(SampleNameSplits_def, AnmailD_DF.loc[:, 'SampleName']))
        AnmailD_DF['Sequence'] = [i[1] for i in SampleNameSplits]
        AnmailD_DF['AnmailID'] = [i[3] for i in SampleNameSplits]
        AnmailD_DF['Days'] = [''.join(i[5:7]).replace('Day', 'D') for i in SampleNameSplits]  # 增加Days列替换 Day 为 D
        AnmailD_DF['Times'] = ['Predose' if i[7] in ['0h'] else i[7] for i in SampleNameSplits]  # 替换0h 为 Predose
        AnmailD_DF['DataAnmailTime'] = AnmailD_DF['Days'] + '-' + AnmailD_DF['AnmailID'] + '-' + AnmailD_DF[
            'Times']  # 组装列数据
        AnmailD_DF_Mean = AnmailD_DF[['runId', 'DataAnmailTime', 'ISArea']].copy()
        STD_QC_cnnectdf = pd.concat([self._get_SDT()[0], self._get_QC()[0]]).copy()  # 合并QC 和STD 用于计算接下来的值
        AnmailD_STDQC = STD_QC_cnnectdf[STD_QC_cnnectdf.loc[:, 'CSRRR_T_F'] == True]  # 筛选STD and QC
        InLableMeanDict = {k: v['ISArea'].mean() for k, v in AnmailD_STDQC.groupby('runId')}
        AnmailD_DF_Mean = AnmailD_DF_Mean[
            AnmailD_DF_Mean.loc[:, 'runId'].isin(list(set(InLableMeanDict.keys())))]  # 筛选runid 在STDQC的df中
        AnmailD_DF_Mean['mean'] = [InLableMeanDict.get(i) for i in AnmailD_DF_Mean['runId']]  # 添加辅助计算列
        AnmailD_DF_Mean['Diff'] = (AnmailD_DF_Mean['ISArea'] - AnmailD_DF_Mean['mean']) / AnmailD_DF_Mean[
            'mean'] * 100  # 计算diff
        AnmailD_DF_Mean['CRROL_T_F'] = [True if -20 <= i <= 20 else False for i in AnmailD_DF_Mean['Diff']]
        # 最大值和最小值的 diffdf 如果是有存在 false 就不参与计算 先筛选出来True 的进行计算
        DiffDF = pd.DataFrame([(name, diffdf[diffdf['CRROL_T_F'] == True]['Diff'].min(),
                                diffdf[diffdf['CRROL_T_F'] == True]['Diff'].max()) for name, diffdf in
                               AnmailD_DF_Mean.groupby('runId')], columns=['runID', 'Min %', 'Max %'])
        return AnmailD_DF_Mean, DiffDF

    def _get_ARA(self):
        '''返回适应性 三个个DF'''
        ARARawDataDF_SAy = self.ARARawDataDF_c[self.ARARawDataDF_c.loc[:, 'SampleKind'].isin(
            ['System Suitability', 'SST'])].copy()  # 晒寻system suitablility
        SAyLdf = []
        for runidname, SAYDF in ARARawDataDF_SAy.groupby('runId'):
            IndexLen = len(SAYDF.index)
            IndexLen_3DF = SAYDF.iloc[IndexLen - 3:, :]  # 筛许最后三针
            SAYDF = IndexLen_3DF[['Area', 'ISArea', 'RetTime', 'ISRetTime']].copy()  # 整理需要的数据
            SAYDF.insert(2, 'A_Ratio_ISA', SAYDF['Area'] / SAYDF['ISArea'])  # 插入比值列
            SAYDF['RetTime'] = SAYDF['RetTime'].astype('float')
            SAYDF['ISRetTime'] = SAYDF['ISRetTime'].astype('float')
            ARICV = SAYDF['A_Ratio_ISA'].std(ddof=1) / SAYDF['A_Ratio_ISA'].mean() * 100
            RTCV = SAYDF['RetTime'].std(ddof=1) / SAYDF['RetTime'].mean() * 100
            ISRTCV = SAYDF['ISRetTime'].std(ddof=1) / SAYDF['ISRetTime'].mean() * 100
            SAyLdf.append([runidname, ARICV, RTCV, ISRTCV])
        SAyLdf = pd.DataFrame(SAyLdf, columns=['runid', 'ARICV %', 'RTCV %', 'ISRTCV %'])
        SAyLdf.set_index('runid', inplace=True)
        SAyLdf_T_F = SAyLdf.applymap(lambda x: True if x < 20 else False)  # 求取真假 小于20%
        return ARARawDataDF_SAy, SAyLdf, SAyLdf_T_F

    def _get_Winnonlin(self):
        '''winnonlin数据raw'''
        Conmuns_W = ['Dosing_Day', 'Group', 'Subject', 'Gender', 'Time', 'Concentration']
        WinDF = self.ARARawDataDF_c[self.ARARawDataDF_c.loc[:, 'SampleKind'] == 'UNKNOWN'].copy()  # 筛选 Unknown
        SampleNameSplits = list(map(SampleNameSplits_def, WinDF.loc[:, 'SampleName']))
        WinDF['Sequence'] = [i[1] for i in SampleNameSplits]
        WinDF['AnmailID'] = [i[3] for i in SampleNameSplits]
        WinDF['Days'] = [i[6] for i in SampleNameSplits]  # 增加Days列替换 Day 为 D
        WinDF['Dosing_Day'] = WinDF['Days']
        WinDF['Group'] = ['G' + i[0] for i in WinDF['AnmailID']]
        WinDF['Subject'] = WinDF['AnmailID']
        WinDF['Gender'] = ['Male' if i[1] == 'M' else 'Female' for i in WinDF['AnmailID']]
        WinDF['Time'] = [float(i[7][0:-1]) / 60 if i[7][-1].lower() == 'm' else i[7][0:-1] for i in
                         SampleNameSplits]  # 转换所有数据到小时
        WinDF['Time'] = WinDF['Time'].astype('float64')
        WinDF['Concentration'] = ['Missing' if pd.isna(i) else i for i in WinDF['Concentration']]
        WinDFresult = WinDF[Conmuns_W]
        return WinDFresult  # 筛选非0 时间点的数据

    def __del__(self):
        pass


def DF_s_ToExcel(sa, AnalyteDescription, RangeName):
    """
    以下为筛选代码或许会用到
    writer = pd.ExcelWriter(abc, engine='openpyxl')
    bbs.StudyID_RawData_All.to_excel(writer, sheet_name='xxx',)
    ws = writer.sheets['xxx']
    ws.auto_filter.ref = "A:P"
    ws.auto_filter.add_filter_column(13, ["Cefepime", "Apple", "Mango"])
    writer.save()
    """
    dds = '_' + datetime.now().strftime('%Y%m%d-%H%M%S')
    excel_file = out_path + '{}-{}-{}.xlsx'.format(sa.StudyName + dds, AnalyteDescription, RangeName)
    writer = pd.ExcelWriter(excel_file)
    OUT_RAW_data = sa.rawdata  # 原始数据
    OUT_STDQCDATA = sa.StudyID_CONCENTRATION  # STDQC原始数据
    OUT_SYX_RAW = sa._get_ARA()[0]  # 适应性原始
    OUT_SYX_value = sa._get_ARA()[1]  # 适应性值
    OUT_SYX_result = sa._get_ARA()[2]  # 适应性判断结果
    OUT_R2_result = sa._get_RunR2()  # r2 结果
    OUT_STD_value = sa._get_SDT()[0]  # STD 原始数据
    OUT_STD_result = sa._get_SDT()[1]  # STD 结果
    OUT_QC_value = sa._get_QC()[0]
    OUT_QC_result = sa._get_QC()[1]
    OUT_LOQ = sa._get_LOQ()[0]  # 残留
    OUT_LOQ_result = sa._get_LOQ()[1]  # 残留
    OUT_Winnonlin = sa._get_Winnonlin()  # Winnonlin
    OUT_Diff_value = sa._get_Diff()[0]  # 内标值
    OUT_Diff_result = sa._get_Diff()[1]  # 内标结果
    OUT_RAW_data.to_excel(writer, sheet_name='RAWDATA')
    OUT_STDQCDATA.to_excel(writer, sheet_name='STDQCDATA')
    OUT_SYX_RAW.to_excel(writer, sheet_name='适应性原始')  # 适应性原始
    OUT_SYX_value.to_excel(writer, sheet_name='适应性值 ')
    OUT_SYX_result.to_excel(writer, sheet_name='适应性判断结果')
    OUT_R2_result.to_excel(writer, sheet_name='R2结果')
    OUT_STD_value.to_excel(writer, sheet_name='STD_RAW')
    OUT_STD_result.to_excel(writer, sheet_name='STD结果')
    OUT_QC_value.to_excel(writer, sheet_name='QC值')
    OUT_QC_result.to_excel(writer, sheet_name='QC结果')
    OUT_LOQ.to_excel(writer, sheet_name='残留')
    OUT_LOQ_result.to_excel(writer, sheet_name='残留结果')
    OUT_Diff_value.to_excel(writer, sheet_name='内标值')
    OUT_Diff_result.to_excel(writer, sheet_name='内标结果')
    CRROLTFDF = OUT_Diff_value['CRROL_T_F'].to_frame()  # 拼接不需要的
    OUT_Winnonlin = pd.merge(OUT_Winnonlin, CRROLTFDF, left_index=True, right_index=True)
    OUT_Winnonlin['analyte'] = AnalyteDescription
    # OUT_Winnonlin.to_excel(writer, sheet_name='Winnonlin')
    writer.save()
    return OUT_Winnonlin


def set_row_height(row):
    '''word 增加行高度'''
    # https://stackoverflow.com/questions/37532283/python-how-to-adjust-row-height-of-table-in-docx
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), "452.830189")
    trHeight.set(qn('w:hRule'), "atLeast")
    trPr.append(trHeight)


class Out_Word_All:
    '''输出内容到word表格'''

    def __init__(self, sa, AnalyteDescription, RangeName, low_limit):
        self.Wmb_word = r'K:\mashuaifei\watson\moban\Watsonmb.docx'
        self.doc = Document(self.Wmb_word)
        self.low_limit = low_limit
        self.sa = sa
        self.AnalyteDescription = AnalyteDescription
        self.RangeName = RangeName
        # self.out_ARA()
        # self.out_STD()
        # self.out_QC()
        # self.out_Diff()
        # self.out_LOQ()
        self.out_winnonlin()
        self.save()

    def write(self, table, tablename, name):
        lie = ['Run Id'] + self.get_lie()
        table.loc[-1] = table.columns  # 增加一行
        table.index = table.index + 1  # 把index的每一项增加1
        table = table.sort_index()  # 重新排序一下
        p = self.doc.add_paragraph()  # 加一段文字
        p.add_run("{}".format(tablename))  # noqa # bold=加粗
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 将该段落居中
        tableSTD1 = self.doc.add_table(table.shape[0], table.shape[1], style=name)
        tbllook = tableSTD1._tblPr.xpath('w:tblLook')[0]
        tbllook.set(qn('w:lastRow'), '1')
        if name == 'STDTableS1':
            for i in range(1, table.shape[0], 2):
                tableSTD1.cell(i, 0).merge(tableSTD1.cell(i + 1, 0))
            for WR in range(table.shape[0]):
                set_row_height(tableSTD1.rows[WR])
                tableSTD1.cell(WR, 0).text = lie[WR]
                tableSTD1.cell(WR, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for WR in range(table.shape[0]):
                set_row_height(tableSTD1.rows[WR])
                for WC in range(1, table.shape[1]):
                    tableSTD1.cell(WR, WC).text = str(table.values[WR][WC])
                    tableSTD1.cell(WR, WC).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            for WR in range(table.shape[0]):
                set_row_height(tableSTD1.rows[WR])
                for WC in range(0, table.shape[1]):
                    tableSTD1.cell(WR, WC).text = str(table.values[WR][WC])
                    tableSTD1.cell(WR, WC).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def out_ARA(self):
        table1 = self.sa._get_ARA()[0].loc[:, ['SampleName']]
        table = self.sa._get_ARA()[1].applymap(lambda x: f'{x:.2f}%')
        table = table.reset_index()
        id = list(set(list(table1['SampleName'].apply(lambda x: x.split(' ')[2]))))
        table['runid'] = table['runid'].apply(lambda x: id[0] + '_Run' + str(x))
        table.rename(columns={'runid': 'Run Id'}, inplace=True)

        table.loc[-1] = table.columns  # 增加一行
        table.index = table.index + 1  # 把index的每一项增加1
        table = table.sort_index()  # 重新排序一下

        rows = table.index.size  # table表的总行数计算行数
        cols = table.columns.size  # table表的总列数计算列数
        p = self.doc.add_paragraph()  # 加一段文字
        p.add_run("{}".format("附表2  系统适应性"))  # noqa # bold=加粗
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 将该段落居中

        tableARA = self.doc.add_table(rows, cols, style='ARATableS')  # 新增加一个表格3行7列并指定表格样式为'tablestylerock' 自定义的表格样式
        tbllook = tableARA._tblPr.xpath('w:tblLook')[0].set(qn('w:lastColumn'), '1')
        for WR in range(rows):
            set_row_height(tableARA.rows[WR])  # 增加行高度
            for WC in range(cols):
                # 插入数据
                tableARA.rows[WR].cells[WC].text = str(table.iloc[WR][WC])
                tableARA.cell(WR, WC).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.doc.add_page_break()

    def get_lie(self):
        table = self.sa._get_SDT()[0].loc[:, ['runId', 'assayDescription']]
        id = list(set(list(table['assayDescription'])))
        table.rename(columns={'runId': 'Run Id'}, inplace=True)
        table['Run Id'] = table['Run Id'].apply(lambda x: id[0] + '_Run' + ("{:.0f} ").format(x))
        id_list = list(set(list(table['Run Id'])))
        id_list.sort(key=list(table['Run Id']).index)

        id_list = [val for val in id_list for i in range(2)]
        return id_list

    def out_STD(self):
        table1 = self.sa._get_RunR2().loc[:, ['RunID', 'R2']]
        table1.rename(columns={'RunID': 'Run Id'}, inplace=True)
        table2 = self.sa._get_SDT()[0].loc[:, ['runId', 'Concentration', 'assayDescription', 'LEVELNUMBER',
                                               'CONCENTRATION', 'C_Ratio_ROL']]
        table2.rename(columns={'runId': 'Run Id'}, inplace=True)
        table4 = self.sa._get_SDT()[1]
        table2['double'] = [2 if i % 2 == 1 else 1 for i in range(table2.shape[0])]
        id = list(set(list(table2['assayDescription'])))
        # table2['runId'] = table2[['runId', 'assayDescription']].apply(lambda x: x['assayDescription'] + '_Run' + str(x['runId']), axis=1)
        table2_1 = table2.pivot_table(index=['Run Id', 'double'], columns='CONCENTRATION',
                                      values='Concentration').applymap(lambda x: f'{x:.2f}')
        table2_2 = table2.pivot_table(index=['Run Id', 'double'], columns='CONCENTRATION',
                                      values='C_Ratio_ROL').applymap(lambda x: f'{x:.2f}')
        table2_1 = table2_1.merge(table1, how='right', on='Run Id')
        table2_1['Run Id'] = table2_1['Run Id'].apply(lambda x: id[0] + '_Run' + str(x))
        table2_1['Run Id'] = [list(table2_1['Run Id'])[i] if i % 2 == 0 else '' for i in range(table2_1.shape[0])]
        table2_1['R2'] = table2_1['R2'].apply(lambda x: str('%.4f' % x))
        table2_1['R2'] = [list(table2_1['R2'])[i] if i % 2 == 0 else '' for i in range(table2_1.shape[0])]
        table2_2 = table2_2.reset_index()
        table2_2.drop(['double'], axis=1, inplace=True)
        self.write(table2_1, '附表3-1', 'STDTableS1')
        self.doc.add_page_break()
        self.write(table2_2, '附表3-2', 'STDTableS1')
        self.doc.add_page_break()

    def out_QC(self):
        table = self.sa._get_QC()[0].loc[:, ['runId', 'Concentration', 'assayDescription', 'LEVELNUMBER',
                                             'CONCENTRATION', 'C_Ratio_ROL']]
        id = list(set(list(table['assayDescription'])))
        table.rename(columns={'runId': 'Run Id'}, inplace=True)
        table['double'] = [2 if i % 2 == 1 else 1 for i in range(table.shape[0])]
        # print(table)
        table['Concentration'] = table['Concentration'].apply(pd.to_numeric, errors='coerce')
        table['C_Ratio_ROL'] = table['C_Ratio_ROL'].apply(pd.to_numeric, errors='coerce')
        table_1 = table.pivot_table(index=['Run Id', 'double'], columns='CONCENTRATION',
                                    values='Concentration').applymap(lambda x: f'{x:.2f}')
        table_2 = table.pivot_table(index=['Run Id', 'double'], columns='CONCENTRATION', values='C_Ratio_ROL').applymap(
            lambda x: f'{x:.2f}')
        table_1 = table_1.reset_index()
        table_2 = table_2.reset_index()
        table_1.drop(['double'], axis=1, inplace=True)
        table_2.drop(['double'], axis=1, inplace=True)

        self.write(table_1, '附表4-1', 'STDTableS1')
        self.doc.add_page_break()
        self.write(table_2, '附表4-2', 'STDTableS1')
        self.doc.add_page_break()

    def out_Diff(self):
        table = self.sa._get_Diff()[1]
        table1 = self.sa._get_SDT()[0].loc[:, ['assayDescription']]
        id = list(set(list(table1['assayDescription'])))
        table.rename(columns={'runID': 'Run Id'}, inplace=True)
        table['Run Id'] = table['Run Id'].apply(lambda x: id[0] + '_Run' + ("{:.0f} ").format(x))
        for i in range(1, table.shape[1]):
            table.iloc[:, i:i + 1] = table.iloc[:, i:i + 1].applymap(
                lambda x: ("{:.2f} ").format(x) if x != '' else '').astype(str)

        self.write(table, '附表5', 'ARATableS')
        self.doc.add_page_break()

    def out_LOQ(self):
        table = self.sa._get_LOQ()[0].loc[:, ['runId', 'SampleName', 'Area', 'ISArea']]
        table.rename(columns={'runId': 'Run Id'}, inplace=True)
        table = table.reset_index(drop=True)
        table['assayDescription'] = table['SampleName'].apply(lambda x: x.split(' ')[2])
        id = list(set(list(table['assayDescription'])))
        table['SampleName'] = ['Carryover-1' if i % 2 == 0 else 'Carryover-2' for i in range(table.shape[0])]
        table['Area'] = table['Area'].apply(lambda x: 'BLOQ' if x != '' else x)
        table['ISArea'] = table['ISArea'].apply(lambda x: 'BLOQ' if x != '' else x)
        table['Run Id'] = table['Run Id'].apply(lambda x: id[0] + '_Run' + str(x))
        table['Run Id'] = [list(table['Run Id'])[i] if i % 2 == 0 else '' for i in range(table.shape[0])]
        table.drop(['assayDescription'], axis=1, inplace=True)

        self.write(table, '附表6', 'STDTableS1')
        self.doc.add_page_break()

    def out_winnonlin(self):

        def del_with(df):
            df_F = df.drop(['Group', 'Dosing_Day', 'Gender'], axis=1)
            df_F = df_F.set_index(['Time', 'Subject'], drop=True)
            df_F = df_F.unstack()
            df_F.columns = list(df_F.columns.levels[1])
            df_F = df_F.reset_index()
            Time_list = []
            for i in list(df_F['Time']):
                if i == 0:
                    Time_list.append(0)
                if 0 < i < 0.5:
                    Time_list.append(str('%.2f' % i))
                if i == 0.5:
                    Time_list.append(0.5)
                if 0.5 < i < 0.1:
                    Time_list.append(str('%.2f' % i))
                if 1 <= i:
                    Time_list.append(str('%.0f' % i))
            df_F['Time'] = Time_list
            mean_columns = df_F[list(df_F.columns)[1:]]
            df_F['Mean'] = mean_columns.mean(axis=1).apply(lambda x: f'{x:.2f}')  ## nan 也可以计算
            sd_columns = df_F[list(df_F.columns)[1:-1]]
            df_F['SD'] = sd_columns.std(axis=1).apply(lambda x: f'{x:.2f}')
            df_F.loc[-1] = df_F.columns  # 增加一行
            df_F.index = df_F.index + 1  # 把index的每一项增加1
            df_F = df_F.sort_index()  # 重新排序一下
            df_F.columns = [str(i) for i in range(df_F.shape[1])]
            df_F = df_F.replace(0, 'BLOQ')
            return df_F

        """获得试验时间顺序"""
        table = self.sa._get_Winnonlin()
        table = table.reset_index(drop=True)
        time = list(set(list(table['Dosing_Day'])))

        time_int = [int(i) for i in time]
        time_int.sort()
        time_int = [str(i) for i in time_int]
        time.sort(key=list(time_int).index)
        """修改数据"""
        table['Concentration'] = table['Concentration'].apply(
            lambda x: 0 if x == 'Missing' or x < self.low_limit else float(str('%.2f' % x)))
        # for i in range(len(time)):
        #     table1 = table[table.Dosing_Day.str.contains(time[i])]
        #     print(table1)
        """根据雌雄，组别，时间分类"""
        table1 = table[table.Dosing_Day.str.contains(time[0])]
        group1 = list(set(list(table1['Group'])))
        group1.sort(key=list(table1['Group']).index)
        for i in group1:
            df = table1[table1.Group.str.contains(i)]
            df_1 = df[df.Gender.str.contains('Female')]
            df_2 = df[df.Gender.str.contains('Male')]
            df_F = del_with(df_1)
            df_M = del_with(df_2)
            df_F = pd.concat([df_F, df_M])
            print(df_F)
            print('<------------------------------------------------------------------------->')
        table2 = table[table.Dosing_Day.str.contains(time[-1])]
        group2 = list(set(list(table2['Group'])))
        group2.sort(key=list(table2['Group']).index)
        print(group2)
        for i in group2:
            df = table2[table2.Group.str.contains(i)]
            df_1 = df[df.Gender.str.contains('Female')]
            df_2 = df[df.Gender.str.contains('Male')]
            df_F = del_with(df_1)
            df_M = del_with(df_2)
            df_F = pd.concat([df_F, df_M])
            print(df_F)
            print('<------------------------------------------------------------------------->')

    def save(self):
        dds = '_' + datetime.now().strftime('%Y%m%d-%H%M%S')
        word_file = out_path + '{}-{}-{}.docx'.format(self.sa.StudyName + dds, self.AnalyteDescription, self.RangeName)
        self.doc.save(word_file)  # 另存为新文件


def watson_excel_word_out(StudyNO, scrange_Dicts):
    '''导出excel和word的整合函数合并了winnonlin'''
    wlsa = WatsonLoadSQL(StudyNO)
    Analytenamesets = set(wlsa._Query_AnalyteDescription()['AnalyteDescription'])  # AnalyteIndex 名字的集合 {'SBK001'}
    # print(Analytenamesets)  查询出药物的名字
    wol = []
    for Ana in Analytenamesets:
        for rangename, v in scrange_Dicts.items():
            # 遍历字典，不明的查询参数 是两个range，rangename对应的是字典的key
            sa = SystemAdaptability(StudyName=StudyNO,
                                    StudyID_Runs_R2=wlsa.StudyID_Runs_R2,
                                    StudyID_RawData_All=wlsa.StudyID_RawData_All,
                                    StudyID_CONCENTRATION=wlsa.StudyID_CONCENTRATION,
                                    Query_AnalyteDescription=wlsa.Query_AnalyteDescription,
                                    AnalyteDescription=Ana,
                                    SDT_Range=v['scrange'],
                                    QC_Range=v['sqarange'])
            low_limit = v['scrange'][0]  # int
            # print(sa)  sa这里还是数据库的查询
            #  这三个参数下面调用了两次
            Out_Word_All(sa, Ana, rangename, low_limit)  # 这一步 把数据写入了word
    #         wol.append(DF_s_ToExcel(sa, Ana, ranger'name))
    # winnonlindf = pd.concat(wol)
    # dds = '_' + datetime.now().strftime('%Y%m%d-%H%M%S')
    # excel_file = out_path + '{}-{}.xlsx'.format(StudyNO + dds, 'winnonlin')
    # excel_file_g_mean = out_path + '{}-{}.xlsx'.format(StudyNO + dds, 'winnonlin-g-mean')
    # winnonlindf = winnonlindf[['analyte', 'Dosing_Day', 'Group', 'Subject', 'Gender', 'Time', 'Concentration']]
    # winnonlindf.sort_values(['analyte', 'Dosing_Day', 'Group', 'Gender', 'Subject', 'Time'], inplace=True)
    # winnonlindf.drop_duplicates(inplace=True)
    # windflist = []
    # for name, df in winnonlindf.groupby(['analyte', 'Dosing_Day', 'Group', 'Subject']):
    #     # 如果某个动物的time 排序后的第一个数据列Concentration 是 missing 那么修改为0
    #     df.iloc[0:1, :].replace('Missing', 0, inplace=True)
    #     windflist.append(df)
    # winnonlindf = pd.concat(windflist)
    # # 求出winnolin的分组后均值数据某些情况下使用.
    # win_g_mean = winnonlindf.groupby(['analyte', 'Dosing_Day', 'Group', 'Gender', 'Time']).apply(
    #     lambda d: d['Concentration'].replace({'Missing': 0}).mean()).to_frame('Concentration').reset_index()
    # winnonlindf.to_excel(excel_file)
    # win_g_mean.to_excel(excel_file_g_mean)
    # return winnonlindf


StudyName = 'A2020010-K10-01'
scrange_Dicts = {'A': {'scrange': [10, 20, 100, 200, 1000, 2000, 8000, 10000], 'sqarange': [30, 500, 7500]}}
watson_excel_word_out(StudyName, scrange_Dicts)
